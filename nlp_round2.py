import pandas as pd
import numpy as npy
import bs4
import requests
import csv
import matplotlib.pyplot as pPlot
import docx
from docx import Document
import operator
import re
import itertools
from collections import Counter
import nltk
from nltk.corpus import wordnet 
from nltk.stem import WordNetLemmatizer
from nltk.stem.porter import PorterStemmer
from nltk.tokenize import sent_tokenize, word_tokenize
import spacy
from spacy import displacy
from spacy import tokenizer
from spacy.matcher import Matcher 
from spacy.tokens import Span 
import urllib.request 
from bs4 import BeautifulSoup
from tqdm import tqdm
from sklearn.metrics import confusion_matrix
from sklearn import metrics
from operator import truediv
from sklearn.metrics import accuracy_score
from sklearn.metrics import precision_score
from sklearn.metrics import recall_score
from sklearn.metrics import f1_score
import networkx as nx
import os
import spacy.cli
# spacy.cli.download("en_core_web_lg") 



from wordcloud import WordCloud, STOPWORDS

doc = Document('cs1.docx')

n=len(doc.tables)
for i in range(0,n):
   doc.tables[0]._element.getparent().remove(doc.tables[0]._element)     

# #docs after removing table is stored as t1.docs

doc.save('t1.docx')

t2 =Document('t1.docx')
doc_1 = docx.Document()
for para in t2.paragraphs :
   new_text = para.text
   temp = re.sub(r'Fig. [0-9][0-9].[0-9]',' ' , new_text)
   temp1 = re.sub(r'Fig. [0-9].[0-9]',' ' , temp)
   temp2=re.sub(r'Fig. [0-9].[0-9][0-9]',' ' , temp1)
   temp3=re.sub(r'Fig. [0-9][0-9].[0-9][0-9]',' ' , temp2)
   temp4=re.sub(r'Table [0-9].[0-9]',' ' , temp3)
   temp5=re.sub(r'Table [0-9][0-9].[0-9]',' ' , temp4)
   temp6=re.sub(r'Sects. [0-9].[0-9]',' ' , temp5)
   temp7=re.sub(r'Sects. [0-9][0-9].[0-9]',' ' , temp6)
   temp8=re.sub(r'[0-9][0-9].[0-9]',' ' , temp7)
   temp9=re.sub(r'[0-9].[0-9]',' ' , temp8)
   temp10=re.sub(r'Chap. [0-9]',' ' , temp9)
   temp11 = re.sub(r"[^a-zA-Z0-9]", " ", temp10)
   temp12=re.sub(r'Chap. [0-9][0-9]',' ' , temp11)
   if len(temp12)>20 :
      doc_1.add_paragraph(temp12)

doc_1.save('t3.docx')

# docs after removing images from t1 is stored as t3.docs

t4=Document('t3.docx')

# tockenisation
tokens = []
for para in t4.paragraphs:
   nltk_tokens = nltk.word_tokenize(para.text,"english",False)
   tokens=tokens+nltk_tokens


tokens_final=[]
for ele in tokens : 
   if ele ==',' or ele=='.' or ele=='(' or ele==')' or ele=='=' or ele==';' or ele=='1' or ele== '+' or ele=='2' or ele=='[' or ele==']' :
      tokens.remove(ele)

#lemmatisation

wordnet_lemmatizer = WordNetLemmatizer()
lema=[]
for w in tokens:
   lema.append(wordnet_lemmatizer.lemmatize(w))


# steming
# nltk.download('wordnet')
stem=[]
porter_stemmer = PorterStemmer()
for w in lema:
   stem.append(porter_stemmer.stem(w))

for ele1 in stem :
   if ele1==':' :
      stem.remove(ele1)

# nltk.download('wordnet')

lex_names={}
for word in stem:
   syn=wordnet.synsets(word)
   for t in syn :
      lex_names[word]=t.lexname()


print(lex_names)

nounDict={}
verbDict={}

for ele,t in lex_names.items():
   if("noun." in t) :
      if(t in nounDict) :
         nounDict[t]+=1
      else :
         nounDict[t]=1
   elif("verb." in t) :
      if(t in verbDict) :
         verbDict[t]+=1
      else :
         verbDict[t]=1

print(nounDict)
print(verbDict)

pPlot.barh(list(nounDict.keys()), nounDict.values(), color='b')
pPlot.show()
pPlot.barh(list(verbDict.keys()), verbDict.values(), color='b')
pPlot.show()

# # part 1 completed  ****************************************************************************************************************************

nlp = spacy.load('en_core_web_sm')

entities=[]
for para in t4.paragraphs:
   text= para.text
   doc = nlp(text)
   ents = [(e.text, e.start_char, e.end_char, e.label_) for e in doc.ents]
   entities.append(ents)
   print(entities)

predicted = []


t5=Document('firstthree.docx')
for para in t5.paragraphs:
   text= para.text
   doc = nlp(text)
   for e in doc.ents :
      ents = [(e.text, e.start_char, e.end_char, e.label_)]
      predicted.append(e.label_)

# fp2=open("pred.txt",'w',encoding='utf-8')
# fp2.writelines('\n'.join(predicted))
# fp2.write('\n')

# with open('actual.txt', 'r' , encoding='utf-8') as infile:
#     true_values = [ i for i in infile]
# with open('pred.txt', 'r' , encoding='utf-8') as infile:
#     predictions = [ i for i in infile]


fileobj=open("pred.txt")
predicted_1=[]
for line in fileobj:
    predicted_1.append(line.strip())
# print(predicted_1)

fileobj=open("actual.txt")
actual_1=[]
for line in fileobj:
    actual_1.append(line.strip())
# print(actual_1)

confusion = confusion_matrix(actual_1, predicted_1 , labels=['ORG' ,'PERSON' ,'CARDINAL' ,'GPE' ,'NORP' ,'DATE' ,'ORDINAL' ,'TIME' ])
print(confusion)
 
precision = precision_score(actual_1, predicted_1,average='micro')
f = f1_score(actual_1, predicted_1,average='micro')
accuracy=accuracy_score(actual_1,predicted_1)

print( precision)
print(f)
print(accuracy)

# part 2 completed *************************************************************************************************************************


def get_relation(sent):
   doc = nlp(sent)
   # Matcher class object
   matcher = Matcher(nlp.vocab)
   relation=[]
   # define the pattern
   pattern = [{'DEP': 'ROOT'},
            {'DEP': 'prep', 'OP': "?"},
            {'DEP': 'agent', 'OP': "?"},
            {'POS': 'ADJ', 'OP': "?"}]
   matcher.add("matching_1", [pattern], on_match=None)
   matches = matcher(doc)
   for mathc_id, start, end in matches:
      matched_span = doc[start: end]
      # print(matched_span.text)
      relation.append(matched_span.text)
   return relation

ar=[]
fp2=open("relation.txt",'w',encoding='utf-8')
for para in t4.paragraphs:
   text= para.text
   a=get_relation(text)
   ar.append(a)
   fp2.writelines('\n'.join(a))



def ne_rel(filename):
  '''Used to preprocess the file for Named entity relationship and detects relationship
  between PERSON entities.'''

  f_s=open(filename)
  remove_chap="[cC]hapter [0-9]+"
  book=""
  for line in f_s:
      book+=line
  book=book.replace("\n"," ")
  book=book.replace('“','"')
  book=book.replace('”','"')
  book=book.replace('—'," ")
  book=book.replace('_'," ")
  book=re.sub(remove_chap,'',book)
  book = re.sub(' +',' ',book)  #For handling multiple spaces

  ent_name=[]
  ent_start=[]
  ent_end=[]
  sent_list=sent_tokenize(book)
  
  mdl=spacy.load("en_core_web_lg")
  labl=mdl(book)
  for e in labl.ents:
    #print(e,"\t",e.label_)
    if e.label_=='PERSON' or e.label_=='ORG':
      ent_name.append(e.text)
      ent_start.append(e.start_char)
      ent_end.append(e.end_char)


  lis=list(zip(ent_name,ent_start,ent_end))
  lis.sort(key=lambda x:x[1])
  #print(lis)
  pat_list=[r'.\b[Bb]elongs\b.',r'.\b[Dd]epends\b.',r'.\b[Dd]enoted\b.',r'.\b[Cc]omputed\b.',r'.\b[Ss]wapped\b.',r'.\b[Aa]ppears\b.']
  type_list=["Belong",'Depends','Denoted','Computed','Swapped','Appears']
  rel_list=list(zip(pat_list,type_list))
  #print(rel_list)
  for i,tup1 in enumerate(lis[:-1]):
    tup2=lis[i+1]
    if tup2[1]-tup1[2]<=100 and str(tup1[0])!=str(tup2[0]):

      sel_text=book[tup1[1]:tup2[2]+1]
      for rel in rel_list:
        reg="\\b"+tup1[0]+"\\b"+rel[0]+"\\b"+tup2[0]+"\\b"
        if re.match(reg,sel_text):
          print(rel[1]+" REL between",tup1[0],"and",tup2[0]+":")
          print(sel_text)
          print("\n")
          pass

ne_rel('textfile.txt')

